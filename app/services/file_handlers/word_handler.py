"""Word 文件处理器 (.docx)

使用 python-docx 进行文本提取，这是 Python 生态中处理 .docx 文件的标准库。

注意：仅支持 .docx 格式（Office 2007+），不支持旧版 .doc 格式（Office 97-2003）。
如需支持 .doc 格式，需要额外安装 python-pptx 或使用 LibreOffice 转换。
"""

from pathlib import Path
from typing import List

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from loguru import logger

from app.config import config
from app.services.file_handlers.base import BaseFileHandler


class WordHandler(BaseFileHandler):
    """Word 文件处理器 (.docx)

    文本提取策略：
    - 使用 python-docx 逐段提取文本
    - 每个段落之间保留换行符以维持结构
    - 表格内容单独处理：每个单元格内容提取后用制表符连接同行单元格，
      不同行之间用换行分隔，保留表格的二维结构
    - 页眉页脚、图片、公式等非文本元素会被忽略

    分片策略（单级分片）：
    - 使用 RecursiveCharacterTextSplitter 按自然段落边界切分
    - chunk_size = 1600 字符（config.chunk_max_size * 2）
    - chunk_overlap = 100 字符
    - 分隔符优先级：双换行 → 单换行 → 空格 → 字符级
    """

    @property
    def supported_extensions(self) -> List[str]:
        return ["docx"]

    def __init__(self):
        self._splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.chunk_max_size * 2,
            chunk_overlap=config.chunk_overlap,
            length_function=len,
            is_separator_regex=False,
        )

    def extract_text(self, file_path: str) -> str:
        """使用 python-docx 从 .docx 中提取文本（含段落和表格）"""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError(
                "python-docx 未安装，请运行: pip install python-docx"
            )

        path = Path(file_path)
        text_parts: List[str] = []

        try:
            doc = DocxDocument(str(path))

            # 使用 iter_block_items 按文档顺序遍历段落和表格
            for block in self._iter_block_items(doc):
                text_parts.append(block)

            paragraph_count = len(doc.paragraphs)
            table_count = len(doc.tables)
            logger.info(
                f"Word 文本提取完成: {file_path}, "
                f"{paragraph_count} 个段落, {table_count} 个表格"
            )

        except Exception as e:
            logger.error(f"Word 文件打开失败: {file_path}, 错误: {e}")
            raise RuntimeError(f"无法读取 Word 文件: {e}") from e

        if not text_parts:
            logger.warning(f"Word 文档中未提取到任何文本: {file_path}")
            return ""

        return "\n\n".join(text_parts)

    def split(self, content: str, file_path: str = "") -> List[Document]:
        if not content or not content.strip():
            logger.warning(f"Word 文档内容为空: {file_path}")
            return []

        docs = self._splitter.create_documents(
            texts=[content],
            metadatas=[{
                "_source": file_path,
                "_extension": ".docx",
                "_file_name": Path(file_path).name,
            }],
        )

        logger.info(f"Word 分割完成: {file_path} -> {len(docs)} 个分片")
        return docs

    @staticmethod
    def _iter_block_items(doc) -> List[str]:
        """
        按文档顺序遍历段落和表格，返回文本块列表。

        python-docx 的 paragraph 和 table 对象是分开存储的，
        此方法根据 XML 元素在文档体中的出现顺序将它们交织在一起，
        确保输出文本的顺序与原文一致。

        表格处理：
        - 每个单元格文本用 \\t 分隔（同行）
        - 每行用换行分隔
        - 不同表格之间保留空行分隔
        """
        from docx.oxml.ns import qn
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        body = doc.element.body
        blocks: List[str] = []

        for child in body.iterchildren():
            if isinstance(child, CT_P):
                para = Paragraph(child, doc)
                text = para.text.strip()
                if text:
                    blocks.append(text)

            elif isinstance(child, CT_Tbl):
                table = Table(child, doc)
                table_text = WordHandler._extract_table_text(table)
                if table_text:
                    blocks.append(table_text)

        return blocks

    @staticmethod
    def _extract_table_text(table) -> str:
        """提取表格文本，保留二维结构"""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            # 过滤全空行
            if any(cells):
                rows.append("\t".join(cells))

        return "\n".join(rows) if rows else ""
